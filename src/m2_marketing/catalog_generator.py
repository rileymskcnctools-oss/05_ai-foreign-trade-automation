"""
FT Workspace v2.0 - M2: Marketing Material Generator

Generate product catalogs, single product sheets, and market-localized versions
as PDF, DOCX, and PPTX files.

Usage:
    from src.m2_marketing.catalog_generator import CatalogGenerator
    gen = CatalogGenerator()
    result = gen.generate_catalog(["GF-001", "GR-001"], output_path="output/catalogs/test.pdf")
"""

import os
import json
from datetime import datetime
from typing import Optional

from fpdf import FPDF
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


# Windows system fonts
WIN_FONT_DIR = os.path.join(os.environ.get("WINDIR", r"C:\Windows"), "Fonts")

def _get_font_path(filename: str) -> str:
    """Get full path to a Windows system font."""
    path = os.path.join(WIN_FONT_DIR, filename)
    if os.path.exists(path):
        return path
    return filename  # fallback: let FPDF search


class CatalogPDF(FPDF):
    """Custom FPDF class for catalog generation."""

    def __init__(self, company_name: str = "FT Workspace", orientation: str = "P"):
        super().__init__(orientation=orientation, unit="mm", format="A4")
        self.company_name = company_name
        self.set_auto_page_break(auto=True, margin=20)
        self.add_font("Arial", "", _get_font_path("arial.ttf"), uni=True)
        self.add_font("Arial", "B", _get_font_path("arialbd.ttf"), uni=True)
        self.add_font("Arial", "I", _get_font_path("ariali.ttf"), uni=True)

    def header(self):
        if self.page_no() == 1:
            return  # No header on first page (has title page)
        self.set_font("Arial", "B", 9)
        self.set_text_color(100, 100, 100)
        self.cell(0, 8, self.company_name, align="L")
        self.cell(0, 8, f"Page {self.page_no()}", align="R", new_x="LMARGIN", new_y="NEXT")
        self.line(10, 12, 200, 12)
        self.ln(4)

    def footer(self):
        self.set_y(-15)
        self.set_font("Arial", "I", 7)
        self.set_text_color(150, 150, 150)
        self.cell(0, 10, f"Generated: {datetime.now().strftime('%Y-%m-%d')} | {self.company_name}", align="C")

    def title_page(self, title: str, subtitle: str = "", date_str: str = ""):
        """Generate a title page."""
        self.add_page()
        self.ln(50)
        self.set_font("Arial", "B", 30)
        self.set_text_color(30, 60, 120)
        self.cell(0, 15, title, align="C", new_x="LMARGIN", new_y="NEXT")
        self.ln(5)
        if subtitle:
            self.set_font("Arial", "", 16)
            self.set_text_color(100, 100, 100)
            self.cell(0, 10, subtitle, align="C", new_x="LMARGIN", new_y="NEXT")
        self.ln(10)
        self.set_font("Arial", "", 12)
        self.set_text_color(130, 130, 130)
        if date_str:
            self.cell(0, 8, date_str, align="C", new_x="LMARGIN", new_y="NEXT")
        self.cell(0, 8, self.company_name, align="C", new_x="LMARGIN", new_y="NEXT")

    def section_header(self, title: str):
        """Section header with colored bar."""
        self.ln(4)
        self.set_fill_color(30, 60, 120)
        self.set_text_color(255, 255, 255)
        self.set_font("Arial", "B", 13)
        self.cell(0, 10, f"  {title}", fill=True, new_x="LMARGIN", new_y="NEXT")
        self.ln(2)

    def body_text(self, text: str, bold: bool = False, size: int = 10):
        """Regular body text."""
        style = "B" if bold else ""
        self.set_font("Arial", style, size)
        self.set_text_color(40, 40, 40)
        self.multi_cell(0, 6, text)
        self.ln(1)

    def bullet_point(self, text: str, indent: int = 10):
        """Bullet point item."""
        self.set_font("Arial", "", 10)
        self.set_text_color(60, 60, 60)
        x = self.get_x()
        self.cell(indent)
        self.cell(5, 6, chr(8226))  # bullet character
        self.multi_cell(0, 6, text)
        self.ln(0.5)

    def spec_table(self, specs: list[tuple]):
        """Two-column specification table."""
        self.ln(2)
        col_w1 = 70
        col_w2 = 110
        self.set_font("Arial", "", 9)
        for i, (label, value) in enumerate(specs):
            if i % 2 == 0:
                self.set_fill_color(240, 245, 255)
            else:
                self.set_fill_color(255, 255, 255)
            self.set_text_color(60, 60, 60)
            self.set_font("Arial", "B", 9)
            self.cell(col_w1, 7, f"  {label}", fill=True)
            self.set_font("Arial", "", 9)
            self.cell(col_w2, 7, f"  {str(value) or 'N/A'}", fill=True, new_x="LMARGIN", new_y="NEXT")
        self.ln(2)

    def product_page(self, product: dict):
        """Generate a single product page."""
        self.add_page()

        # Product name header
        self.set_font("Arial", "B", 20)
        self.set_text_color(30, 60, 120)
        self.cell(0, 12, product.get("product_name_en", "N/A"), new_x="LMARGIN", new_y="NEXT")

        # Product code and category
        self.set_font("Arial", "", 11)
        self.set_text_color(100, 100, 100)
        code = product.get("product_code", "N/A")
        cat = product.get("category", "")
        sub = product.get("sub_category", "")
        self.cell(0, 7, f"Code: {code}  |  {cat} / {sub}", new_x="LMARGIN", new_y="NEXT")
        self.ln(3)

        # Specifications
        self.section_header("Specifications")
        specs = []
        for field, label in [
            ("material", "Material"),
            ("handle_material", "Handle"),
            ("length_cm", "Length (cm)"),
            ("weight_kg", "Weight (kg)"),
            ("head_width_cm", "Head Width (cm)"),
            ("hardness", "Hardness"),
            ("surface_treatment", "Surface Treatment"),
            ("moq", "MOQ (pcs)"),
            ("packaging_type", "Packaging"),
            ("qty_per_carton", "Qty/Carton"),
            ("carton_size_cm", "Carton Size (cm)"),
            ("gw_per_carton_kg", "GW/Carton (kg)"),
            ("lead_time_days", "Lead Time (days)"),
            ("loading_qty_20ft", "Loading 20'"),
            ("loading_qty_40ft", "Loading 40'"),
            ("loading_qty_40hq", "Loading 40HQ"),
        ]:
            val = product.get(field)
            if val is not None and str(val).strip():
                specs.append((label, val))
        self.spec_table(specs)

        # Selling angle
        angle = product.get("selling_angle", "")
        if angle:
            self.section_header("Key Features")
            self.body_text(angle)

        # Use scenario
        scenario = product.get("use_scenario", "")
        if scenario:
            self.section_header("Applications")
            self.body_text(scenario)

        # Keywords
        keywords = product.get("target_keywords", "")
        if keywords:
            self.ln(2)
            self.set_font("Arial", "I", 8)
            self.set_text_color(130, 130, 130)
            self.cell(0, 6, f"Keywords: {keywords}")


class CatalogGenerator:
    """Generate product catalogs and marketing materials."""

    def __init__(self, company_name: str = "FT Workspace", db=None):
        self.company_name = company_name
        self.db = db
        if db is None:
            from src.core.database import FTDatabase
            self.db = FTDatabase()

    def generate_catalog(
        self,
        product_codes: list[str] = None,
        category: str = None,
        output_path: str = None,
        market: str = None,
        format: str = "pdf",
    ) -> dict:
        """
        Generate a multi-product catalog.

        Args:
            product_codes: List of product codes. If None and category is set, use all in category.
            category: Filter by category.
            output_path: Output file path. Auto-generated if None.
            market: Target market (for localization).
            format: Output format (pdf, docx).

        Returns:
            Result dict with filepath, product_count, errors.
        """
        result = {"filepath": None, "product_count": 0, "errors": []}

        # Get products
        if product_codes:
            products = []
            for code in product_codes:
                p = self.db.product_get(code)
                if p:
                    products.append(p)
                else:
                    result["errors"].append(f"Product not found: {code}")
        elif category:
            products = self.db.product_list(category=category, status=None, limit=1000)
        else:
            products = self.db.product_list(status=None, limit=1000)

        if not products:
            result["errors"].append("No products found for the given criteria.")
            return result

        result["product_count"] = len(products)

        # Generate output path
        if output_path is None:
            ts = datetime.now().strftime("%Y%m%d_%H%M%S")
            cat_label = category or "all"
            output_path = os.path.join(
                "output", "catalogs",
                f"catalog_{cat_label.lower().replace(' ', '_')}_{ts}.{format}"
            )

        os.makedirs(os.path.dirname(output_path), exist_ok=True)

        try:
            if format == "pdf":
                self._generate_pdf_catalog(products, output_path, market)
            elif format == "docx":
                self._generate_docx_catalog(products, output_path, market)
            else:
                result["errors"].append(f"Unsupported format: {format}")
                return result
        except Exception as e:
            result["errors"].append(str(e))
            return result

        result["filepath"] = output_path
        return result

    def generate_single_sheet(
        self,
        product_code: str,
        output_path: str = None,
        format: str = "pdf",
    ) -> dict:
        """Generate a single product sheet (A4 page)."""
        result = {"filepath": None, "errors": []}

        product = self.db.product_get(product_code)
        if not product:
            result["errors"].append(f"Product not found: {product_code}")
            return result

        if output_path is None:
            ts = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = os.path.join(
                "output", "catalogs",
                f"sheet_{product_code}_{ts}.{format}"
            )

        os.makedirs(os.path.dirname(output_path), exist_ok=True)

        try:
            if format == "pdf":
                pdf = CatalogPDF(self.company_name)
                pdf.product_page(product)
                pdf.output(output_path)
            elif format == "docx":
                doc = Document()
                self._docx_product_page(doc, product)
                doc.save(output_path)
            else:
                result["errors"].append(f"Unsupported format: {format}")
                return result
        except Exception as e:
            result["errors"].append(str(e))
            return result

        result["filepath"] = output_path
        return result

    def _generate_pdf_catalog(self, products: list[dict], output_path: str, market: str = None):
        """Generate PDF catalog."""
        pdf = CatalogPDF(self.company_name)

        # Title page
        title = "Product Catalog"
        subtitle = ""
        if market:
            subtitle = f"Market Edition: {market}"
        pdf.title_page(title, subtitle, datetime.now().strftime("%B %Y"))

        # Table of contents
        pdf.add_page()
        pdf.section_header("Table of Contents")
        pdf.ln(2)
        pdf.set_font("Arial", "", 11)
        pdf.set_text_color(40, 40, 40)
        for i, p in enumerate(products):
            code = p.get("product_code", "N/A")
            name = p.get("product_name_en", "N/A")
            pdf.cell(0, 8, f"  {i+1}. {code} - {name}", new_x="LMARGIN", new_y="NEXT")

        # Product pages
        for product in products:
            pdf.product_page(product)

        pdf.output(output_path)

    def _generate_docx_catalog(self, products: list[dict], output_path: str, market: str = None):
        """Generate DOCX catalog."""
        doc = Document()

        # Style setup
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        # Title
        title = doc.add_heading('Product Catalog', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if market:
            doc.add_paragraph(f'Market Edition: {market}', style='Subtitle')
        doc.add_paragraph(datetime.now().strftime('%B %Y')).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        # Products
        for product in products:
            self._docx_product_page(doc, product)
            doc.add_page_break()

        doc.save(output_path)

    def _docx_product_page(self, doc: Document, product: dict):
        """Add a product section to a DOCX document."""
        # Product name
        name = product.get("product_name_en", "N/A")
        code = product.get("product_code", "N/A")
        h = doc.add_heading(name, level=2)
        for run in h.runs:
            run.font.color.rgb = RGBColor(30, 60, 120)

        # Code and category
        p = doc.add_paragraph()
        cat = product.get("category", "")
        sub = product.get("sub_category", "")
        run = p.add_run(f"Code: {code}  |  {cat} / {sub}")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(100, 100, 100)

        # Specifications table
        doc.add_heading("Specifications", level=3)
        table = doc.add_table(rows=0, cols=2, style="Light Shading Accent 1")
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        spec_fields = [
            ("material", "Material"),
            ("handle_material", "Handle"),
            ("length_cm", "Length (cm)"),
            ("weight_kg", "Weight (kg)"),
            ("head_width_cm", "Head Width (cm)"),
            ("hardness", "Hardness"),
            ("surface_treatment", "Surface Treatment"),
            ("moq", "MOQ (pcs)"),
            ("packaging_type", "Packaging"),
            ("qty_per_carton", "Qty/Carton"),
            ("carton_size_cm", "Carton Size (cm)"),
            ("gw_per_carton_kg", "GW/Carton (kg)"),
            ("lead_time_days", "Lead Time (days)"),
            ("loading_qty_20ft", "Loading 20'"),
            ("loading_qty_40ft", "Loading 40'"),
            ("loading_qty_40hq", "Loading 40HQ"),
        ]

        for field, label in spec_fields:
            val = product.get(field)
            if val is not None and str(val).strip():
                row = table.add_row()
                row.cells[0].text = label
                row.cells[1].text = str(val)
                # Set column widths
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph.style.font.size = Pt(9)

        # Key Features
        angle = product.get("selling_angle", "")
        if angle:
            doc.add_heading("Key Features", level=3)
            doc.add_paragraph(angle)

        # Applications
        scenario = product.get("use_scenario", "")
        if scenario:
            doc.add_heading("Applications", level=3)
            doc.add_paragraph(scenario)
